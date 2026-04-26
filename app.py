import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import sys
import copy
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────
# PLACEHOLDER REPLACEMENT (handles runs inside paragraphs
# and table cells, including split-run placeholders)
# ─────────────────────────────────────────────────────

def _replace_in_paragraph(para, mapping):
    """Substitui placeholders preservando negrito/itálico/fonte de cada run."""
    # Passo 1: substituição direta dentro de cada run (preserva formatação 100%)
    for run in para.runs:
        for key, val in mapping.items():
            if key in run.text:
                run.text = run.text.replace(key, val)

    # Passo 2: placeholders divididos entre runs (cirúrgico — só mescla os runs afetados)
    for key, val in mapping.items():
        while True:
            texts = [r.text for r in para.runs]
            combined = "".join(texts)
            if key not in combined:
                break
            # Verifica se ainda está dividido (passo 1 já tratou os inteiros)
            if any(key in t for t in texts):
                break  # já foi resolvido no passo 1, não deveria chegar aqui

            # Acha os índices de runs que contêm o placeholder dividido
            start_pos = combined.index(key)
            end_pos   = start_pos + len(key)
            pos = 0
            start_run = end_run = -1
            for i, t in enumerate(texts):
                if start_run == -1 and pos + len(t) > start_pos:
                    start_run = i
                if pos + len(t) >= end_pos:
                    end_run = i
                    break
                pos += len(t)

            if start_run == -1 or end_run == -1 or start_run == end_run:
                break  # não encontrado ou já no mesmo run

            # Mescla só os runs que fazem parte do placeholder
            merged = "".join(r.text for r in para.runs[start_run:end_run + 1])
            para.runs[start_run].text = merged.replace(key, val)
            for r in para.runs[start_run + 1:end_run + 1]:
                r.text = ""
            break

def _replace_in_element(element, mapping):
    """Percorre qualquer elemento XML e substitui placeholders em todos os parágrafos."""
    from docx.text.paragraph import Paragraph
    for para_elem in element.iter(qn('w:p')):
        para = Paragraph(para_elem, None)
        _replace_in_paragraph(para, mapping)

def replace_placeholders(doc, mapping):
    """Substitui placeholders em todo o documento, incluindo caixas de texto."""
    # Parágrafos normais do corpo
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, mapping)

    # Caixas de texto no corpo do documento
    for txbx in doc.element.body.iter(qn('w:txbxContent')):
        _replace_in_element(txbx, mapping)

    # Cabeçalhos e rodapés (incluindo caixas de texto dentro deles)
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    _replace_in_paragraph(para, mapping)
                for txbx in hf._element.iter(qn('w:txbxContent')):
                    _replace_in_element(txbx, mapping)

# ─────────────────────────────────────────────────────
# MONTH NAMES (PT-BR)
# ─────────────────────────────────────────────────────
MONTHS = [
    "", "janeiro","fevereiro","março","abril","maio","junho",
    "julho","agosto","setembro","outubro","novembro","dezembro"
]

def format_date_full(date_str):
    """Convert DD/MM/AAAA → DD de mês de AAAA"""
    try:
        d, m, y = date_str.strip().split("/")
        return f"{int(d)} de {MONTHS[int(m)]} de {y}"
    except Exception:
        return date_str

# ─────────────────────────────────────────────────────
# TOOLTIP HELPER
# ─────────────────────────────────────────────────────
class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip = None
        widget.bind("<Enter>", self.show)
        widget.bind("<Leave>", self.hide)

    def show(self, _=None):
        x, y, *_ = self.widget.bbox("insert") if hasattr(self.widget, "bbox") else (0,0,0,0)
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25
        self.tip = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        tk.Label(tw, text=self.text, background="#ffffe0", relief="solid",
                 borderwidth=1, font=("Segoe UI", 9)).pack()

    def hide(self, _=None):
        if self.tip:
            self.tip.destroy()
            self.tip = None

# ─────────────────────────────────────────────────────
# SCROLLABLE FRAME
# ─────────────────────────────────────────────────────
class ScrollFrame(tk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self.canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0)
        self.vsb = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.inner = tk.Frame(self.canvas)
        self.inner.bind("<Configure>", lambda e: self.canvas.configure(
            scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0,0), window=self.inner, anchor="nw")
        self.canvas.configure(yscrollcommand=self.vsb.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.vsb.pack(side="right", fill="y")
        self.canvas.bind("<MouseWheel>", self._on_mw)
        self.inner.bind("<MouseWheel>", self._on_mw)

    def _on_mw(self, e):
        self.canvas.yview_scroll(int(-1*(e.delta/120)), "units")

# ─────────────────────────────────────────────────────
# FIELD BUILDER HELPERS
# ─────────────────────────────────────────────────────
PAD = {"padx": 8, "pady": 4}

def make_field(parent, row, label, var, width=40, tooltip=None):
    lbl = ttk.Label(parent, text=label)
    lbl.grid(row=row, column=0, sticky="e", **PAD)
    ent = ttk.Entry(parent, textvariable=var, width=width)
    ent.grid(row=row, column=1, sticky="w", **PAD)
    if tooltip:
        ToolTip(ent, tooltip)
    return ent

def make_label_section(parent, row, text):
    sep = ttk.Separator(parent, orient="horizontal")
    sep.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(12,2))
    ttk.Label(parent, text=text, font=("Segoe UI", 10, "bold"),
              foreground="#1a5276").grid(row=row+1, column=0, columnspan=2,
              sticky="w", padx=8, pady=(0,4))
    return row+2

# ─────────────────────────────────────────────────────
# MAIN APPLICATION
# ─────────────────────────────────────────────────────
class DiplomaApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("📄 Gerador de Diploma e Histórico")
        self.resizable(True, True)
        self.geometry("820x680")
        self._setup_style()
        self._build_ui()

    # ── STYLE ──────────────────────────────────────
    def _setup_style(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        s.configure("TFrame", background="#f5f6fa")
        s.configure("TLabel", background="#f5f6fa", font=("Segoe UI", 10))
        s.configure("TEntry", font=("Segoe UI", 10))
        s.configure("TNotebook", background="#e8eaf0")
        s.configure("TNotebook.Tab", font=("Segoe UI", 10, "bold"), padding=[12,4])
        s.configure("Header.TLabel", font=("Segoe UI", 13, "bold"),
                    foreground="#154360", background="#d6eaf8")
        s.configure("Action.TButton", font=("Segoe UI", 11, "bold"),
                    padding=10, background="#1a5276", foreground="white")
        s.map("Action.TButton", background=[("active","#1f618d")])
        self.configure(bg="#f5f6fa")

    # ── BUILD UI ───────────────────────────────────
    def _build_ui(self):
        # Header banner
        banner = tk.Frame(self, bg="#154360", pady=12)
        banner.pack(fill="x")
        tk.Label(banner, text="Gerador de Diploma e Histórico Escolar",
                 font=("Segoe UI", 15, "bold"), bg="#154360", fg="white").pack()
        tk.Label(banner, text="Preencha os dados e selecione o modelo Word para gerar o documento",
                 font=("Segoe UI", 9), bg="#154360", fg="#aed6f1").pack()

        # Notebook
        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=10, pady=8)

        # Dados gerais (shared between both docs)
        self.vars_geral = {}
        tab_geral = ttk.Frame(nb)
        nb.add(tab_geral, text="👤 Dados do Aluno")
        self._build_tab_geral(tab_geral)

        # Histórico-specific
        self.vars_hist = {}
        tab_hist = ttk.Frame(nb)
        nb.add(tab_hist, text="📋 Histórico Escolar")
        self._build_tab_historico(tab_hist)

        # Diploma-specific
        self.vars_dip = {}
        tab_dip = ttk.Frame(nb)
        nb.add(tab_dip, text="🎓 Diploma")
        self._build_tab_diploma(tab_dip)

        # Discipline grades
        self.vars_disc = {}
        tab_disc = ttk.Frame(nb)
        nb.add(tab_disc, text="📚 Disciplinas")
        self._build_tab_disciplinas(tab_disc)

        # Bottom bar
        bar = tk.Frame(self, bg="#e8eaf0", pady=8)
        bar.pack(fill="x", side="bottom")
        ttk.Button(bar, text="📋 Gerar Histórico", style="Action.TButton",
                   command=self.gerar_historico).pack(side="left", padx=12)
        ttk.Button(bar, text="🎓 Gerar Diploma", style="Action.TButton",
                   command=self.gerar_diploma).pack(side="left", padx=4)
        ttk.Button(bar, text="📄 Gerar Ambos", style="Action.TButton",
                   command=self.gerar_ambos).pack(side="left", padx=4)
        ttk.Button(bar, text="🗑 Limpar Campos",
                   command=self.limpar_campos).pack(side="right", padx=12)

    # ── TAB DADOS GERAIS ───────────────────────────
    def _build_tab_geral(self, parent):
        sf = ScrollFrame(parent, bg="#f5f6fa")
        sf.pack(fill="both", expand=True)
        f = sf.inner
        f.columnconfigure(1, weight=1)

        V = self.vars_geral
        def sv(k): V[k] = tk.StringVar(); return V[k]

        r = 0
        r = make_label_section(f, r, "Identificação do(a) Aluno(a)")
        make_field(f, r,   "Aluno(a):",            sv("aluno"),          tooltip="Nome completo do aluno"); r+=1
        make_field(f, r,   "Data de Nascimento:",   sv("data_nasc"),      tooltip="Formato: DD/MM/AAAA"); r+=1
        make_field(f, r,   "Nacionalidade:",        sv("nacionalidade")); r+=1
        make_field(f, r,   "Naturalidade:",         sv("naturalidade")); r+=1
        make_field(f, r,   "UF:",                   sv("uf"),             width=8); r+=1
        make_field(f, r,   "Filiação 1 (Pai/Mãe):", sv("filiacao_1"),     width=55); r+=1
        make_field(f, r,   "Filiação 2 (Pai/Mãe):", sv("filiacao_2"),     width=55); r+=1
        make_field(f, r,   "CPF:",                  sv("cpf"),            tooltip="Formato: 000.000.000-00"); r+=1
        make_field(f, r,   "RG:",                   sv("rg")); r+=1
        make_field(f, r,   "Órgão Emissor:",        sv("orgao_emissor")); r+=1

        r = make_label_section(f, r, "Curso Anterior")
        make_field(f, r,   "Curso:",                sv("curso_ant")); r+=1
        make_field(f, r,   "Estabelecimento:",      sv("estab_ant"),      width=55); r+=1
        make_field(f, r,   "Ano:",                  sv("ano_ant"),        width=10); r+=1
        make_field(f, r,   "Cidade:",               sv("cidade_ant")); r+=1

        r = make_label_section(f, r, "Dados do Curso Atual")
        make_field(f, r,   "Turma:",                sv("turma"),          width=20); r+=1
        make_field(f, r,   "Data de Início:",       sv("data_inicio"),    tooltip="Formato: DD/MM/AAAA"); r+=1
        make_field(f, r,   "Data de Término:",      sv("data_termino"),   tooltip="Formato: DD/MM/AAAA"); r+=1
        make_field(f, r,   "Frequência (%):",       sv("frequencia"),     width=10); r+=1
        make_field(f, r,   "Resultado:",            sv("resultado")); r+=1
        make_field(f, r,   "Código SISTEC:",        sv("cod_sistec")); r+=1
        make_field(f, r,   "Código Censo:",         sv("cod_censo")); r+=1
        make_field(f, r,   "Carga Horária Estágio:", sv("carga_estagio"), width=55); r+=1

    # ── TAB HISTÓRICO ─────────────────────────────
    def _build_tab_historico(self, parent):
        sf = ScrollFrame(parent, bg="#f5f6fa")
        sf.pack(fill="both", expand=True)
        f = sf.inner
        f.columnconfigure(1, weight=1)

        V = self.vars_hist
        def sv(k): V[k] = tk.StringVar(); return V[k]

        r = 0
        r = make_label_section(f, r, "Dados do Registro")
        make_field(f, r, "Data do Histórico:", sv("data_hist"),
                   tooltip="Formato: DD/MM/AAAA — será convertida para '01 de janeiro de 2024'"); r+=1

        info = ttk.Label(f, text=(
            "ℹ️  Todos os campos pessoais são preenchidos na aba 'Dados do Aluno'.\n"
            "     Os campos de disciplinas são preenchidos na aba 'Disciplinas'.\n"
            "     Configure os placeholders no seu modelo Word conforme o guia abaixo."),
            font=("Segoe UI", 9), foreground="#555", background="#f5f6fa",
            justify="left", wraplength=580)
        info.grid(row=r, column=0, columnspan=2, padx=10, pady=6, sticky="w"); r+=1

        self._build_placeholder_guide(f, r, "historico")

    # ── TAB DIPLOMA ───────────────────────────────
    def _build_tab_diploma(self, parent):
        sf = ScrollFrame(parent, bg="#f5f6fa")
        sf.pack(fill="both", expand=True)
        f = sf.inner
        f.columnconfigure(1, weight=1)

        V = self.vars_dip
        def sv(k): V[k] = tk.StringVar(); return V[k]

        r = 0
        r = make_label_section(f, r, "Dados do Registro do Diploma")
        make_field(f, r, "Número do Registro:", sv("num_registro"),    width=15); r+=1
        make_field(f, r, "Folha:",              sv("folha"),            width=10); r+=1
        make_field(f, r, "Livro nº:",           sv("livro"),            width=10); r+=1
        make_field(f, r, "Data do Diploma:",    sv("data_diploma"),
                   tooltip="Formato: DD/MM/AAAA — será convertida para '01 de janeiro de 2024'"); r+=1
        make_field(f, r, "Data de Conclusão:",  sv("data_conclusao"),   tooltip="Formato: DD/MM/AAAA"); r+=1
        make_field(f, r, "Nome da Escola:",     sv("nome_escola"),      width=55); r+=1
        make_field(f, r, "Município/UF Escola:", sv("municipio_uf"),    width=30); r+=1
        make_field(f, r, "Expedido em:",        sv("expedido_em"),      tooltip="Data de expedição do RG (DD/MM/AAAA)"); r+=1

        info = ttk.Label(f, text=(
            "ℹ️  Os demais campos (nome, CPF, RG, etc.) vêm da aba 'Dados do Aluno'."),
            font=("Segoe UI", 9), foreground="#555", background="#f5f6fa")
        info.grid(row=r, column=0, columnspan=2, padx=10, pady=6, sticky="w"); r+=1

        self._build_placeholder_guide(f, r, "diploma")

    # ── TAB DISCIPLINAS ───────────────────────────
    def _build_tab_disciplinas(self, parent):
        top = tk.Frame(parent, bg="#f5f6fa", pady=6)
        top.pack(fill="x", padx=10)

        ttk.Label(top, text="Quantidade de disciplinas:").pack(side="left")
        self.num_disc_var = tk.IntVar(value=10)
        sb = ttk.Spinbox(top, from_=1, to=40, textvariable=self.num_disc_var,
                         width=4, command=self._refresh_disciplinas)
        sb.pack(side="left", padx=6)
        ttk.Button(top, text="Atualizar", command=self._refresh_disciplinas).pack(side="left")

        info = tk.Label(parent,
            text="ℹ  O nome da disciplina é apenas referência visual — só a Nota/Situação é inserida no documento  (placeholder {{NOTA_1}}, {{NOTA_2}}…)",
            font=("Segoe UI", 9), fg="#555", bg="#f5f6fa", anchor="w", justify="left")
        info.pack(fill="x", padx=12, pady=(0,4))

        self.disc_scroll = ScrollFrame(parent, bg="#f5f6fa")
        self.disc_scroll.pack(fill="both", expand=True, padx=10, pady=4)
        self.disc_inner = self.disc_scroll.inner
        self.disc_inner.columnconfigure(2, weight=1)
        self._refresh_disciplinas()

    def _refresh_disciplinas(self):
        for w in self.disc_inner.winfo_children():
            w.destroy()
        self.vars_disc.clear()
        n = self.num_disc_var.get()

        # Header
        ttk.Label(self.disc_inner, text="#",  font=("Segoe UI",9,"bold"), width=3).grid(
            row=0, column=0, padx=4, pady=2)
        ttk.Label(self.disc_inner, text="Disciplina (referência)",
                  font=("Segoe UI",9,"bold")).grid(row=0, column=1, sticky="w", padx=4)
        ttk.Label(self.disc_inner, text="Nota / Situação → inserida no doc",
                  font=("Segoe UI",9,"bold"), foreground="#154360").grid(
                  row=0, column=2, sticky="w", padx=4)

        for i in range(n):
            label_v = tk.StringVar()   # só referência visual, não vai pro doc
            nota_v  = tk.StringVar()   # vai pro documento via {{NOTA_X}}
            self.vars_disc[f"nota_{i+1}"] = nota_v
            # guardamos o label pra limpar depois
            self.vars_disc[f"_label_{i+1}"] = label_v

            ttk.Label(self.disc_inner, text=f"{i+1}", width=3, anchor="e").grid(
                row=i+1, column=0, sticky="e", padx=(4,2), pady=2)
            ttk.Entry(self.disc_inner, textvariable=label_v, width=34,
                      foreground="#777").grid(row=i+1, column=1, sticky="w", padx=4, pady=2)
            ttk.Entry(self.disc_inner, textvariable=nota_v, width=18,
                      font=("Segoe UI", 10, "bold")).grid(
                      row=i+1, column=2, sticky="w", padx=4, pady=2)

    # ── PLACEHOLDER GUIDE ─────────────────────────
    def _build_placeholder_guide(self, parent, row, doc_type):
        frame = ttk.LabelFrame(parent, text="📌 Guia de Placeholders para o Modelo Word",
                               padding=10)
        frame.grid(row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=8)
        frame.columnconfigure(0, weight=1)

        if doc_type == "historico":
            placeholders = HISTORICO_PLACEHOLDERS
        else:
            placeholders = DIPLOMA_PLACEHOLDERS

        text_box = tk.Text(frame, height=14, width=72, font=("Courier New", 9),
                           bg="#f0f4f8", relief="flat", wrap="none")
        text_box.pack(fill="both", expand=True)
        scroll = ttk.Scrollbar(frame, command=text_box.yview)
        scroll.pack(side="right", fill="y")
        text_box.configure(yscrollcommand=scroll.set)

        text_box.insert("end", placeholders)
        text_box.configure(state="disabled")
        ttk.Button(frame, text="📋 Copiar",
                   command=lambda: self._copy_text(placeholders)).pack(anchor="e", pady=4)

    def _copy_text(self, text):
        self.clipboard_clear()
        self.clipboard_append(text)
        messagebox.showinfo("Copiado", "Texto copiado! Cole no seu editor de referência.")

    # ── BUILD MAPPING ─────────────────────────────
    def _build_mapping(self):
        g = self.vars_geral
        def g_(k): return g[k].get().strip() if k in g else ""

        mapping = {
            "{{ALUNO}}":            g_("aluno"),
            "{{DATA_NASC}}":        g_("data_nasc"),
            "{{NACIONALIDADE}}":    g_("nacionalidade"),
            "{{NATURALIDADE}}":     g_("naturalidade"),
            "{{UF}}":               g_("uf"),
            "{{FILIACAO_1}}":       g_("filiacao_1"),
            "{{FILIACAO_2}}":       g_("filiacao_2"),
            "{{CPF}}":              g_("cpf"),
            "{{RG}}":               g_("rg"),
            "{{ORGAO_EMISSOR}}":    g_("orgao_emissor"),
            "{{CURSO_ANT}}":        g_("curso_ant"),
            "{{ESTAB_ANT}}":        g_("estab_ant"),
            "{{ANO_ANT}}":          g_("ano_ant"),
            "{{CIDADE_ANT}}":       g_("cidade_ant"),
            "{{TURMA}}":            g_("turma"),
            "{{DATA_INICIO}}":      g_("data_inicio"),
            "{{DATA_TERMINO}}":     g_("data_termino"),
            "{{FREQUENCIA}}":       g_("frequencia"),
            "{{RESULTADO}}":        g_("resultado"),
            "{{COD_SISTEC}}":       g_("cod_sistec"),
            "{{COD_CENSO}}":        g_("cod_censo"),
            "{{CARGA_ESTAGIO}}":    g_("carga_estagio"),
            # Naturalidade/UF combinado
            "{{NATURALIDADE_UF}}":  f"{g_('naturalidade')}/{g_('uf')}",
        }

        # Disciplinas — só as notas vão pro documento (nomes são só referência visual)
        for k, v in self.vars_disc.items():
            if k.startswith("_label_"):
                continue
            mapping[f"{{{{{k.upper()}}}}}"] = v.get().strip()

        return mapping

    def _build_mapping_historico(self):
        m = self._build_mapping()
        h = self.vars_hist
        data_hist = h["data_hist"].get().strip() if "data_hist" in h else ""
        m["{{DATA_HIST}}"] = format_date_full(data_hist)
        m["{{DATA_HIST_CURTA}}"] = data_hist
        return m

    def _build_mapping_diploma(self):
        m = self._build_mapping()
        d = self.vars_dip
        def d_(k): return d[k].get().strip() if k in d else ""
        data_dip = d_("data_diploma")
        m["{{NUM_REGISTRO}}"] = d_("num_registro")
        m["{{FOLHA}}"]        = d_("folha")
        m["{{LIVRO}}"]        = d_("livro")
        m["{{DATA_DIPLOMA}}"] = format_date_full(data_dip)
        m["{{DATA_DIPLOMA_CURTA}}"] = data_dip
        m["{{DATA_CONCLUSAO}}"] = d_("data_conclusao")
        m["{{NOME_ESCOLA}}"]  = d_("nome_escola")
        m["{{MUNICIPIO_UF}}"] = d_("municipio_uf")
        m["{{EXPEDIDO_EM}}"]  = d_("expedido_em")
        return m

    # ── GENERATE ──────────────────────────────────
    def _select_template(self, title):
        path = filedialog.askopenfilename(
            title=f"Selecionar Modelo — {title}",
            filetypes=[("Documentos Word", "*.docx"), ("Todos", "*.*")])
        return path or None

    def _select_save(self, default_name):
        path = filedialog.asksaveasfilename(
            title="Salvar como…",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Documentos Word", "*.docx")])
        return path or None

    def gerar_historico(self):
        tpl = self._select_template("Histórico Escolar")
        if not tpl: return
        aluno = self.vars_geral.get("aluno", tk.StringVar()).get().strip()
        out = self._select_save(f"Historico_{aluno or 'aluno'}.docx")
        if not out: return
        try:
            doc = Document(tpl)
            replace_placeholders(doc, self._build_mapping_historico())
            doc.save(out)
            messagebox.showinfo("✅ Sucesso", f"Histórico gerado!\n{out}")
        except Exception as e:
            messagebox.showerror("Erro", str(e))

    def gerar_diploma(self):
        tpl = self._select_template("Diploma")
        if not tpl: return
        aluno = self.vars_geral.get("aluno", tk.StringVar()).get().strip()
        out = self._select_save(f"Diploma_{aluno or 'aluno'}.docx")
        if not out: return
        try:
            doc = Document(tpl)
            replace_placeholders(doc, self._build_mapping_diploma())
            doc.save(out)
            messagebox.showinfo("✅ Sucesso", f"Diploma gerado!\n{out}")
        except Exception as e:
            messagebox.showerror("Erro", str(e))

    def gerar_ambos(self):
        self.gerar_historico()
        self.gerar_diploma()

    def limpar_campos(self):
        for d in [self.vars_geral, self.vars_hist, self.vars_dip, self.vars_disc]:
            for v in d.values():
                v.set("")


# ─────────────────────────────────────────────────────
# PLACEHOLDER GUIDES
# ─────────────────────────────────────────────────────
HISTORICO_PLACEHOLDERS = """\
PLACEHOLDERS PARA O MODELO DO HISTÓRICO ESCOLAR
Copie e cole estes marcadores exatamente no seu documento Word
onde o valor variável deve aparecer:

=== DADOS DO ALUNO ===
{{ALUNO}}             → Nome completo do(a) aluno(a)
{{DATA_NASC}}         → Data de nascimento (DD/MM/AAAA)
{{NACIONALIDADE}}     → Nacionalidade
{{NATURALIDADE}}      → Naturalidade (cidade)
{{UF}}                → UF de naturalidade
{{FILIACAO_1}}        → Filiação 1 (1º responsável)
{{FILIACAO_2}}        → Filiação 2 (2º responsável)
{{CPF}}               → CPF
{{RG}}                → RG
{{ORGAO_EMISSOR}}     → Órgão emissor do RG

=== CURSO ANTERIOR ===
{{CURSO_ANT}}         → Curso anterior
{{ESTAB_ANT}}         → Estabelecimento
{{ANO_ANT}}           → Ano de conclusão
{{CIDADE_ANT}}        → Cidade

=== DADOS DO CURSO ===
{{TURMA}}             → Turma
{{DATA_INICIO}}       → Data de início
{{DATA_TERMINO}}      → Data de término
{{FREQUENCIA}}        → Frequência (%)
{{RESULTADO}}         → Resultado
{{COD_SISTEC}}        → Código SISTEC
{{COD_CENSO}}         → Código Censo
{{CARGA_ESTAGIO}}     → Carga horária de estágio

=== DATA ===
{{DATA_HIST}}         → Ex.: 1 de março de 2024 (por extenso)
{{DATA_HIST_CURTA}}   → Ex.: 01/03/2024

=== NOTAS DAS DISCIPLINAS ===
As disciplinas já estão no modelo — coloque apenas o
placeholder da nota na célula/campo correspondente:

{{NOTA_1}}            → Nota/Situação da 1ª disciplina
{{NOTA_2}}            → Nota/Situação da 2ª disciplina
{{NOTA_3}}            → Nota/Situação da 3ª disciplina
... (até {{NOTA_40}})
"""

DIPLOMA_PLACEHOLDERS = """\
PLACEHOLDERS PARA O MODELO DO DIPLOMA
Copie e cole estes marcadores exatamente no seu documento Word:

=== DADOS DO ALUNO ===
{{ALUNO}}             → Nome completo
{{NATURALIDADE_UF}}   → Cidade/UF  (ex.: Rio de Janeiro/RJ)
{{NATURALIDADE}}      → Apenas a cidade
{{UF}}                → Apenas a UF
{{DATA_NASC}}         → Data de nascimento
{{RG}}                → Número do RG
{{ORGAO_EMISSOR}}     → Órgão emissor do RG
{{EXPEDIDO_EM}}       → Data de expedição do RG
{{CPF}}               → CPF

=== CURSO ANTERIOR ===
{{CURSO_ANT}}         → Ex.: ENSINO MÉDIO
{{ANO_ANT}}           → Ano de conclusão
{{ESTAB_ANT}}         → Estabelecimento

=== REGISTRO ===
{{NUM_REGISTRO}}      → Número do registro
{{FOLHA}}             → Folha
{{LIVRO}}             → Livro nº
{{NOME_ESCOLA}}       → Nome da escola emissora
{{MUNICIPIO_UF}}      → Município e UF da escola

=== DATAS ===
{{DATA_CONCLUSAO}}    → Data de conclusão do curso (DD/MM/AAAA)
{{DATA_DIPLOMA}}      → Ex.: 01 de março de 2024 (por extenso)
{{DATA_DIPLOMA_CURTA}} → Ex.: 01/03/2024

=== OUTROS ===
{{COD_SISTEC}}        → Código SISTEC
{{COD_CENSO}}         → Código Censo
{{CARGA_ESTAGIO}}     → Carga horária de estágio
"""

# ─────────────────────────────────────────────────────
if __name__ == "__main__":
    app = DiplomaApp()
    app.mainloop()
